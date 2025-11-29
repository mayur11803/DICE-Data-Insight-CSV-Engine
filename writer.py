from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from datetime import datetime

def report_writer(df, stats, summary, charts):
    doc = Document()

    # Cover Page
    doc.add_paragraph().add_run("\n" * 3)
    title = doc.add_paragraph()
    title_run = title.add_run("CSV Insight Report")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Date: {datetime.today().strftime('%B %d, %Y')}")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Statistics Overview", level=1)
    doc.add_paragraph(stats)

    # Placeholder for charts
    doc.add_heading("Generated Charts", level=1)
    if charts:
        for path in charts:
            doc.add_picture(path, width=Inches(6))
    else:
        doc.add_paragraph("No charts generated.")

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer  # ✅ required for app.py to call `.getvalue()`

